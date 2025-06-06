import os
import uuid
import json
from typing import Dict, Any
from datetime import datetime
from .base import BaseAgent
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import google.generativeai as genai
import logging

logger = logging.getLogger(__name__)

class FileOutputAgent(BaseAgent):
    """Agent that creates different document types based on text input and document type configuration"""
    
    def __init__(self, config: Dict[str, Any], system_instruction: str = ""):
        super().__init__(config, system_instruction)
        
        # Set up output directory
        self.output_dir = "_OUTPUT"
        os.makedirs(self.output_dir, exist_ok=True)
        
        # Get document type from config
        self.document_type = config.get("document_type", "txt")
        self.use_ai_formatting = config.get("use_ai_formatting", True)
        
        # Initialize Gemini if API key is provided and AI formatting is enabled
        self.model = None
        if self.use_ai_formatting and config.get("api_key"):
            try:
                genai.configure(api_key=config.get("api_key"))
                self.model = genai.GenerativeModel(config.get("model_name", "gemini-1.5-flash"))
            except Exception as e:
                logger.warning(f"Failed to initialize Gemini for AI formatting: {e}")
                self.use_ai_formatting = False
    
    async def process(self, input_data: Any, context: Dict[str, Any] = None) -> Dict[str, Any]:
        """Process input text and create a file of the specified document type"""
        context = context or {}
        
        try:
            # Extract text input
            if isinstance(input_data, dict):
                if "text" in input_data:
                    input_text = input_data["text"]
                elif "output" in input_data:
                    input_text = input_data["output"]
                elif "content" in input_data:
                    input_text = input_data["content"]
                else:
                    input_text = str(input_data)
            else:
                input_text = str(input_data)
            
            if not input_text or input_text.strip() == "":
                return {
                    "error": "No text content provided",
                    "output": "Error: No text content provided",
                    "agent_type": "file_output"
                }
            
            # Generate formatted content using AI if enabled
            formatted_content = input_text
            if self.use_ai_formatting and self.model:
                formatted_content = await self._format_content_with_ai(input_text)
            
            # Create file based on document type
            file_path = await self._create_file(formatted_content)
            
            # Extract filename and format information
            full_filename = os.path.basename(file_path)
            filename_without_ext = os.path.splitext(full_filename)[0]
            file_format = os.path.splitext(full_filename)[1]
            
            return {
                "output": f"{full_filename}",
                "file_path": file_path,
                "filename": filename_without_ext,
                "file_format": file_format,
                "document_type": self.document_type,
                "agent_type": "file_output",
                "token_usage": {"total_tokens": 0}
            }
            
        except Exception as e:
            logger.error(f"Error in FileOutputAgent: {str(e)}")
            return {
                "error": str(e),
                "output": f"Error creating file: {str(e)}",
                "agent_type": "file_output",
                "token_usage": {"total_tokens": 0}
            }
    
    def _clean_markdown_code_blocks(self, content: str) -> str:
        """Remove markdown code block syntax from content"""
        lines = content.strip().split('\n')
        
        # Check if content starts with markdown code block (```python, ```java, ```js, etc.)
        if lines and lines[0].strip().startswith('```'):
            # Remove the first line (```language)
            lines = lines[1:]
        
        # Check if content ends with markdown code block
        if lines and lines[-1].strip() == '```':
            # Remove the last line (```)
            lines = lines[:-1]
        
        # Handle case where there might be extra empty lines after removing backticks
        while lines and not lines[0].strip():
            lines = lines[1:]
        while lines and not lines[-1].strip():
            lines = lines[:-1]
        
        # Join back and strip any extra whitespace
        cleaned_content = '\n'.join(lines).strip()
        return cleaned_content
    
    async def _format_content_with_ai(self, content: str) -> str:
        """Use AI to format content based on document type and content analysis"""
        try:
            formatting_prompt = self._get_formatting_prompt(content)
            
            response = self.model.generate_content(formatting_prompt)
            formatted_content = response.text.strip()
            
            # Clean markdown code blocks if present and if it's code-related content
            if self.document_type in ["py", "java", "c", "cpp", "js", "ts"]:
                formatted_content = self._clean_markdown_code_blocks(formatted_content)
            
            return formatted_content
            
        except Exception as e:
            logger.warning(f"AI formatting failed, using original content: {e}")
            return content
    
    def _get_formatting_prompt(self, content: str) -> str:
        """Generate appropriate formatting prompt based on document type"""
        base_prompt = f"""You are an expert document formatter. Format the following content appropriately for a {self.document_type.upper()} document.

Document Type: {self.document_type}
Original Content: {content}

Instructions:
"""
        
        if self.document_type == "docx":
            return base_prompt + """
- Structure the content with appropriate headings and sections
- If it appears to be a CV/resume, format it with proper sections (Contact Info, Summary, Experience, Education, Skills, etc.)
- If it's a report, use proper headings, subheadings, and paragraph structure
- If it's code documentation, organize it with clear sections and examples
- Use proper paragraph spacing and organization
- Return the content in a structured format with clear headings marked by ## for main sections and ### for subsections
"""
        
        elif self.document_type == "pdf":
            return base_prompt + """
- Structure the content for professional PDF file.
- Use clear headings and proper document hierarchy
- If it's a CV, format with professional sections
- If it's a story or article, use proper narrative structure
- If it's a report, use executive summary, main content, and conclusions
- If it's documentation, organize it with clear sections and examples
- If it's a technical document, include code blocks and explanations
- Ensure proper paragraph breaks and readability
- Return structured content with headings marked by ## and ###
"""
        
        elif self.document_type in ["py", "java", "c", "cpp", "js", "ts"]:
            return base_prompt + f"""
- Format as proper {self.document_type} code
- Add appropriate comments and documentation
- Include proper imports/includes if needed
- Use correct syntax and indentation
- If it's a class, include proper structure
- If it's functions, include docstrings/comments
- Ensure the code is well-organized and follows best practices
- Return only the formatted code without markdown backticks
"""
        
        else:  # txt or other formats
            return base_prompt + """
- Format as plain text with proper structure
- Use clear paragraphs and sections
- Add appropriate spacing and organization
- Ensure readability and logical flow
"""
    
    async def _create_file(self, content: str) -> str:
        """Create the actual file based on document type"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        unique_id = str(uuid.uuid4())[:8]
        
        if self.document_type == "docx":
            return self._create_docx_file(content, timestamp, unique_id)
        elif self.document_type == "pdf":
            return self._create_pdf_file(content, timestamp, unique_id)
        elif self.document_type in ["py", "java", "c", "cpp", "js", "ts", "html", "css"]:
            return self._create_code_file(content, timestamp, unique_id)
        else:  # Default to txt
            return self._create_txt_file(content, timestamp, unique_id)
    
    def _create_docx_file(self, content: str, timestamp: str, unique_id: str) -> str:
        """Create a DOCX file"""
        filename = f"document_{timestamp}_{unique_id}.docx"
        file_path = os.path.join(self.output_dir, filename)
        
        doc = Document()
        
        # Parse content and create structured document
        lines = content.split('\n')
        current_paragraph = []
        
        for line in lines:
            line = line.strip()
            if not line:
                if current_paragraph:
                    doc.add_paragraph(' '.join(current_paragraph))
                    current_paragraph = []
                continue
            
            # Check for headings
            if line.startswith('## '):
                if current_paragraph:
                    doc.add_paragraph(' '.join(current_paragraph))
                    current_paragraph = []
                heading = doc.add_heading(line[3:], level=1)
                heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif line.startswith('### '):
                if current_paragraph:
                    doc.add_paragraph(' '.join(current_paragraph))
                    current_paragraph = []
                heading = doc.add_heading(line[4:], level=2)
                heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                current_paragraph.append(line)
        
        # Add remaining content
        if current_paragraph:
            doc.add_paragraph(' '.join(current_paragraph))
        
        doc.save(file_path)
        return file_path
    
    def _create_pdf_file(self, content: str, timestamp: str, unique_id: str) -> str:
        """Create a PDF file (using reportlab)"""
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.enums import TA_LEFT, TA_CENTER
            
            filename = f"document_{timestamp}_{unique_id}.pdf"
            file_path = os.path.join(self.output_dir, filename)
            
            doc = SimpleDocTemplate(file_path, pagesize=letter,
                                    rightMargin=72, leftMargin=72,
                                    topMargin=72, bottomMargin=18)
            
            styles = getSampleStyleSheet()
            story = []
            
            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                spaceAfter=30,
                alignment=TA_CENTER
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=14,
                spaceAfter=12,
                spaceBefore=12
            )
            
            # Parse content
            lines = content.split('\n')
            current_paragraph = []
            
            for line in lines:
                line = line.strip()
                if not line:
                    if current_paragraph:
                        para = Paragraph(' '.join(current_paragraph), styles['Normal'])
                        story.append(para)
                        story.append(Spacer(1, 12))
                        current_paragraph = []
                    continue
                
                if line.startswith('## '):
                    if current_paragraph:
                        para = Paragraph(' '.join(current_paragraph), styles['Normal'])
                        story.append(para)
                        story.append(Spacer(1, 12))
                        current_paragraph = []
                    heading = Paragraph(line[3:], title_style)
                    story.append(heading)
                elif line.startswith('### '):
                    if current_paragraph:
                        para = Paragraph(' '.join(current_paragraph), styles['Normal'])
                        story.append(para)
                        story.append(Spacer(1, 12))
                        current_paragraph = []
                    heading = Paragraph(line[4:], heading_style)
                    story.append(heading)
                else:
                    current_paragraph.append(line)
            
            # Add remaining content
            if current_paragraph:
                para = Paragraph(' '.join(current_paragraph), styles['Normal'])
                story.append(para)
            
            doc.build(story)
            return file_path
            
        except ImportError:
            # Fallback to text file if reportlab is not available
            logger.warning("reportlab not available, creating text file instead of PDF")
            return self._create_txt_file(content, timestamp, unique_id)
    
    def _create_code_file(self, content: str, timestamp: str, unique_id: str) -> str:
        """Create a code file with appropriate extension"""
        filename = f"code_{timestamp}_{unique_id}.{self.document_type}"
        file_path = os.path.join(self.output_dir, filename)
        
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content)
        
        return file_path
    
    def _create_txt_file(self, content: str, timestamp: str, unique_id: str) -> str:
        """Create a plain text file"""
        filename = f"document_{timestamp}_{unique_id}.txt"
        file_path = os.path.join(self.output_dir, filename)
        
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content)
        
        return file_path
    
    @staticmethod
    def get_supported_document_types() -> Dict[str, str]:
        """Get a list of supported document types"""
        return {
            "txt": "Plain Text",
            "docx": "Microsoft Word Document",
            "pdf": "PDF Document",
            "py": "Python Code",
            "java": "Java Code",
            "c": "C Code",
            "cpp": "C++ Code",
            "js": "JavaScript Code",
            "ts": "TypeScript Code",
            "html": "HTML Document",
            "css": "CSS Stylesheet"
        }
